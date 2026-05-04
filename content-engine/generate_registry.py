
import csv
import json
import re
from pathlib import Path
from docx import Document

def clean_text(text):
    if not text: return ""
    # Remove zero-width spaces and other non-printable chars
    text = text.replace('\u200b', '').replace('\xa0', ' ')
    return text.strip()

def parse_registry():
    registry = {}
    
    base_path = Path(__file__).parent / "Example Resources"
    
    # 1. Parse state list.csv
    csv_path = base_path / "state list.csv"
    if csv_path.exists():
        with open(csv_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                exam_name = clean_text(row.get("Name of the Examination", ""))
                if exam_name:
                    registry[exam_name] = {
                        "state_or_ut": clean_text(row.get("State", "")),
                        "category": clean_text(row.get("Category of Exam", "")),
                        "conducting_body": clean_text(row.get("Conducting Body", "")),
                        "website": clean_text(row.get("Website", ""))
                    }

    # 2. Parse central exams list.docx
    central_path = base_path / "central exams list.docx"
    if central_path.exists():
        doc = Document(central_path)
        for table in doc.tables:
            # Header: Sr.No | Category | Conducting Body | Name of Examination | WEBSITE | Main head
            rows = list(table.rows)
            if not rows: continue
            for row in rows[1:]: # Skip header
                cells = [clean_text(c.text) for c in row.cells]
                if len(cells) >= 4:
                    exam_name = cells[3]
                    if exam_name:
                        registry[exam_name] = {
                            "state_or_ut": "Central",
                            "category": cells[1],
                            "conducting_body": cells[2],
                            "website": cells[4] if len(cells) > 4 else ""
                        }

    # 3. Parse UT Exams EDTD.docx
    ut_path = base_path / "UT Exams EDTD.docx"
    if ut_path.exists():
        doc = Document(ut_path)
        for table in doc.tables:
            # Header: Ser No | State | Category of Exam | Conducting Body | Ser No | Name of the Examination | Website | Logo
            rows = list(table.rows)
            if not rows: continue
            header = [clean_text(c.text) for c in rows[0].cells]
            if "Name of the Examination" in header:
                name_idx = header.index("Name of the Examination")
                state_idx = header.index("State") if "State" in header else -1
                cat_idx = header.index("Category of Exam") if "Category of Exam" in header else -1
                body_idx = header.index("Conducting Body") if "Conducting Body" in header else -1
                web_idx = header.index("Website") if "Website" in header else -1
                
                for row in rows[1:]:
                    cells = [clean_text(c.text) for c in row.cells]
                    if len(cells) > name_idx:
                        exam_name = cells[name_idx]
                        if exam_name:
                            registry[exam_name] = {
                                "state_or_ut": cells[state_idx] if state_idx != -1 else "UT",
                                "category": cells[cat_idx] if cat_idx != -1 else "",
                                "conducting_body": cells[body_idx] if body_idx != -1 else "",
                                "website": cells[web_idx] if web_idx != -1 else ""
                            }

    # Save to file
    output_path = Path("master_exam_registry.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(registry, f, indent=2, ensure_ascii=False)
    
    print(f"Registry generated with {len(registry)} entries.")
    return registry

if __name__ == "__main__":
    parse_registry()
