
import sys
from docx import Document
from pathlib import Path

def read_docx(path):
    print(f"\n--- READING: {path} ---")
    try:
        doc = Document(path)
        # Read first 50 paragraphs and last 50 to get a sense of structure
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        print(f"Total non-empty paragraphs: {len(paras)}")
        if paras:
            print("\nFIRST 20 PARAGRAPHS:")
            for p in paras[:20]:
                print(p)
        
        print(f"\nTotal tables: {len(doc.tables)}")
        for i, table in enumerate(doc.tables[:5]): # Check first 5 tables
            print(f"\nTable {i+1} rows: {len(table.rows)}")
            for j, row in enumerate(table.rows[:50]): # Check first 50 rows
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                print(f"Row {j+1}: {' | '.join(cells)}")

    except Exception as e:
        print(f"Error reading {path}: {e}")

if __name__ == "__main__":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    files = [
        r"k:\H DRIVE\Quantum Climb\APPS\VeerNXT\DocumentEngine\Example Resources\central exams list.docx",
        r"k:\H DRIVE\Quantum Climb\APPS\VeerNXT\DocumentEngine\Example Resources\UT Exams EDTD.docx"
    ]
    for f in files:
        read_docx(f)
