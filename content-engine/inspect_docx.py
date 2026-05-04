#!/usr/bin/env python3
"""Inspect the structure of the test DOCX file."""

from docx import Document

docx_path = r'Example Resources\SSC Stenographer Grade C&D\5. TEST SERIES-10\Section 1.docx'
doc = Document(docx_path)

print('Answer section (lines 757-879):')
for i in range(757, min(880, len(doc.paragraphs))):
    if doc.paragraphs[i].text.strip():
        print(f'{i}: {repr(doc.paragraphs[i].text)}')
